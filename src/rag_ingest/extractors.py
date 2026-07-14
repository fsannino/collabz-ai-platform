"""Extração de texto por formato."""

from __future__ import annotations

import logging
from pathlib import Path

log = logging.getLogger("rag_ingest")

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".html",
    ".htm",
    ".cmake",
}


def extract_text(path: Path) -> str:
    extension = path.suffix.lower()

    try:
        if extension == ".pdf":
            return _from_pdf(path)
        if extension == ".docx":
            return _from_docx(path)
        if extension in {".txt", ".md", ".cmake"}:
            return path.read_text(encoding="utf-8", errors="replace")
        if extension in {".html", ".htm"}:
            return _from_html(path)
    except Exception as error:
        log.warning("Falha extraindo %s: %s", path, error)

    return ""


def _from_pdf(path: Path) -> str:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def _from_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(part for part in parts if part.strip())


def _from_html(path: Path) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(
        path.read_text(encoding="utf-8", errors="replace"),
        "html.parser",
    )
    return soup.get_text(separator="\n")
